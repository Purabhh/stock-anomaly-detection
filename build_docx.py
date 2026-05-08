"""
Build final_report.docx from final_report.md.

Pandoc's default markdown->docx conversion produces tables without visible
borders. This script post-processes the docx so every table cell gets a
solid 0.5pt border, and the header row gets bold text + light-grey shading
to match standard academic-report formatting.

Run after editing final_report.md:
    python build_docx.py

Requires: pip install pypandoc-binary python-docx
"""

import pypandoc
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_border(cell):
    """Apply solid 0.5pt black border to all four edges of a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn('w:tcBorders')):
        tc_pr.remove(existing)
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')         # 0.5pt
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_cell_shading(cell, fill_hex):
    """Set a background fill color on a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def style_tables(doc):
    """Border every cell and emphasize header rows."""
    for i, table in enumerate(doc.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)
        for cell in table.rows[0].cells:
            set_cell_shading(cell, 'D9D9D9')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        print(f'  Table {i}: {len(table.rows)} rows x {len(table.columns)} cols')


def main():
    pypandoc.convert_file('final_report.md', 'docx', outputfile='final_report.docx')
    print('Converted final_report.md -> final_report.docx')

    doc = Document('final_report.docx')
    print(f'Found {len(doc.tables)} tables; styling each:')
    style_tables(doc)
    doc.save('final_report.docx')
    print('Saved final_report.docx with table borders + header shading')


if __name__ == '__main__':
    main()
